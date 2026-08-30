"""Document parsers.

The caller resolves rights and passes bytes it is allowed to use. This module
never fetches a remote URL and never touches the filesystem outside the paths
given to it.
"""

from __future__ import annotations

import io
import logging
import re
import shutil
import subprocess
import tempfile
from collections import defaultdict
from contextlib import contextmanager
from dataclasses import dataclass, field
from statistics import median
from time import perf_counter
from typing import Callable, Dict, List, Optional

from app.core.encoding import RepairPlan, apply_plan, build_plan
from app.core.pdf_encoding import repair_pdf
from app.core.pdf_profile import PdfProfile, is_tagged, measure_page
from app.core.struct_tree import TreeHealth, inspect
from app.core import columns, lexicon, tagged_reader
from app.core.quality import assess

_CONTROL_CHARACTERS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_EXCESS_BLANK_LINES = re.compile(r"\n{3,}")
_TRAILING_SPACES = re.compile(r"[ \t]+\n")

# En dessous, le texte est assez mauvais pour justifier le filet de secours.
_RESCUE_FLOOR = 0.90


class UnsupportedMediaType(ValueError):
    """Le type de fichier envoyé n'est pas pris en charge."""


class OcrUnavailable(RuntimeError):
    """OCRmyPDF n'est pas disponible dans l'image."""


class OcrFailed(RuntimeError):
    """L'OCR a échoué sur ce document."""


# Signatures des formats qu'on sait lire. Trois octets suffisent à trancher,
# et cela referme deux problèmes d'un coup : un dépôt légitime dont le
# navigateur a mal deviné l'extension, et un fichier arbitraire annoncé comme
# PDF qui partirait dans un analyseur qui ne l'attend pas.
_SIGNATURES = (
    (b"%PDF-", "application/pdf"),
    (b"PK\x03\x04", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
)


def sniff(payload: bytes) -> Optional[str]:
    """Le type réel du fichier, lu dans ses premiers octets.

    Renvoie None quand la signature n'est pas reconnue : un fichier texte n'en
    a pas, et l'absence de signature n'est pas une raison de refuser.
    """

    head = payload[:8].lstrip()
    for signature, media_type in _SIGNATURES:
        if head.startswith(signature):
            return media_type
    return None


def normalise(text: str) -> str:
    """Clean extracted text without destroying its structure.

    Blank lines are the signal the chunker uses to find block boundaries, so
    they are collapsed but never removed.
    """

    cleaned = _CONTROL_CHARACTERS.sub("", text.replace("\r\n", "\n").replace("\r", "\n"))
    cleaned = _TRAILING_SPACES.sub("\n", cleaned)
    cleaned = _EXCESS_BLANK_LINES.sub("\n\n", cleaned)
    return cleaned.strip()


def load_plain_text(payload: bytes) -> str:
    return normalise(payload.decode("utf-8", errors="replace"))


logger = logging.getLogger(__name__)


def _table_to_markdown(table) -> str:
    """Rendre un tableau en Markdown, une ligne par ligne du tableau.

    Le programme officiel EST un tableau : semaine, partie, leçon,
    compétences. Le rendre linéairement entrelacerait les colonnes et le
    rendrait illisible, pour un humain comme pour un modèle.
    """

    rows = []
    for row in table:
        cells = [
            " ".join((cell or "").split()) for cell in row
        ]
        if any(cells):
            rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows)


def _clamp(box, page_box):
    """Ramener une boîte dans les limites de sa page."""

    left, top, right, bottom = box
    page_left, page_top, page_right, page_bottom = page_box
    return (
        max(left, page_left),
        max(top, page_top),
        min(right, page_right),
        min(bottom, page_bottom),
    )


def _read_pdf_text(payload: bytes) -> tuple:
    """Lire un PDF en respectant sa mise en page. Renvoie (texte, nb de pages).

    pdfplumber connaît la position de chaque mot : les tableaux sont extraits
    comme tableaux, et le texte hors tableau garde son ordre de lecture.
    """

    try:
        import pdfplumber
    except ImportError as error:  # pragma: no cover
        raise UnsupportedMediaType("pdfplumber n'est pas installé") from error

    pages = []
    with pdfplumber.open(io.BytesIO(payload)) as document:
        page_count = len(document.pages)
        for index, page in enumerate(document.pages, start=1):
            parts = []

            tables = page.find_tables()
            for table in tables:
                rendered = _table_to_markdown(table.extract())
                if rendered:
                    parts.append(rendered)

            # Le texte situé hors des tableaux, dans son ordre de lecture.
            outside = page
            for table in tables:
                # Un tableau peut déborder la page d'une fraction de point,
                # par arrondi. pdfplumber refuse alors la découpe et fait
                # échouer tout le document : on ramène la boîte dans la page.
                try:
                    outside = outside.outside_bbox(_clamp(table.bbox, page.bbox))
                except ValueError:
                    logger.warning("tableau hors page, découpe ignorée")
            remaining = outside.extract_text(layout=False) or ""
            if remaining.strip():
                parts.append(remaining.strip())

            # Le numéro de page devient un titre, donc un repère de
            # citation : un élève doit pouvoir être renvoyé à la bonne page.
            pages.append("\n\n".join(parts) if parts else "")

    rendered = "\n\n".join(
        f"## p. {number}\n\n{body}" for number, body in enumerate(pages, start=1) if body.strip()
    )
    return normalise(rendered), page_count


@dataclass(frozen=True)
class PdfReading:
    """Tout ce qu'une seule lecture du document a permis d'établir.

    Chaque brique — profil, filets, polices, texte — ouvrait auparavant le
    fichier de son côté : quatre ouvertures pour un document, vingt et une
    secondes avant d'avoir lu une ligne sur un PDF de 99 pages. Elles
    partagent désormais la même passe.
    """

    pages: List[str]
    plan: RepairPlan
    fonts: List
    profile: PdfProfile
    ruled_pages: List[int]
    script: str = "latin"
    timings: Dict[str, float] = field(default_factory=dict)
    tree: TreeHealth = field(default_factory=TreeHealth)
    read_by_tree: bool = False


def _read_columns(pages_words: List[list], widths: List[float]) -> List[str]:
    """Rendre le document en suivant ses gouttières de blanc."""

    metrics = columns.measure(pages_words)
    lines = [columns.visual_lines(words, metrics) for words in pages_words]
    known = columns.furniture(lines)
    per_page = [
        columns.gutters(page_lines, width, metrics)
        for page_lines, width in zip(lines, widths)
    ]
    fallback = columns.template(per_page, len(lines), metrics)
    return [
        columns.render(
            page_lines,
            # Le gabarit ne secourt qu'une page qui a vu des colonnes sans les
            # voir toutes. Une page qui n'en voit aucune est de la prose.
            [(left + right) / 2 for left, right in found]
            if len(found) >= len(fallback) or not found
            else fallback,
            width,
            metrics,
            known,
        )
        for page_lines, found, width in zip(lines, per_page, widths)
    ]


@contextmanager
def _step(timings: Dict[str, float], name: str):
    """Chronométrer une étape.

    La fusion des passes fait perdre la visibilité qu'on avait quand chaque
    brique était séparée. Sans ce relevé, la prochaine régression de
    performance serait invisible.
    """

    start = perf_counter()
    try:
        yield
    finally:
        timings[name] = round(timings.get(name, 0.0) + perf_counter() - start, 3)


def read_pdf_document(payload: bytes, repair: bool = True) -> PdfReading:
    """Lire un PDF de bout en bout, en une passe et en se chronométrant.

    L'ordre est contraint : la correction d'encodage doit précéder la lecture,
    puisqu'elle change ce que le lecteur verra. Tout le reste — surface de
    texte, surface d'image, filets, mots, tableaux — se relève ensuite d'un
    seul parcours.
    """

    try:
        import pdfplumber
    except ImportError as error:  # pragma: no cover
        raise UnsupportedMediaType("pdfplumber n'est pas installé") from error

    timings: Dict[str, float] = {}
    fonts: list = []
    script = "latin"

    if repair:
        # On corrige d'abord le PDF lui-même : chaque police reçoit la table
        # qui décode réellement son texte. Le document sort alors juste dès la
        # première lecture, au lieu d'être rattrapé mot à mot ensuite.
        with _step(timings, "encodage"):
            try:
                payload, fonts, script = repair_pdf(payload)
            except Exception:  # noqa: BLE001
                logger.warning("réécriture du PDF impossible, lecture telle quelle")

    pages: List[str] = []
    pages_words: List[list] = []
    widths: List[float] = []
    marks: Dict[tuple, List[str]] = defaultdict(list)
    coverages: List[float] = []
    needing_ocr: List[int] = []
    ruled: List[int] = []

    with _step(timings, "lecture"):
        with pdfplumber.open(io.BytesIO(payload)) as document:
            for number, page in enumerate(document.pages, start=1):
                parts = []
                tables = page.find_tables()
                for table in tables:
                    rendered = _table_to_markdown(table.extract())
                    if rendered:
                        parts.append(rendered)

                outside = page
                for table in tables:
                    # Un tableau peut déborder la page d'une fraction de point,
                    # par arrondi. pdfplumber refuse alors la découpe et fait
                    # échouer tout le document : on ramène la boîte dans la page.
                    try:
                        outside = outside.outside_bbox(_clamp(table.bbox, page.bbox))
                    except ValueError:
                        logger.warning("tableau hors page, découpe ignorée")
                remaining = outside.extract_text(layout=False) or ""
                if remaining.strip():
                    parts.append(remaining.strip())
                pages.append("\n\n".join(parts))

                try:
                    words = page.extract_words(extra_attrs=["fontname", "mcid"])
                except Exception:  # noqa: BLE001
                    # Un PDF peut refuser de livrer ses polices. On garde alors
                    # le texte tel quel : sans police, aucune correction n'est
                    # sûre.
                    logger.warning("polices illisibles sur une page")
                    words = []
                pages_words.append(words)

                # Le texte de chaque contenu marqué, relevé au passage : c'est
                # le pont vers l'arbre de structure, et il ne coûte rien ici.
                for word in words:
                    mark = word.get("mcid")
                    if mark is not None:
                        marks[(number, int(mark))].append(word["text"])

                widths.append(float(page.width))
                measure_page(number, page, words, coverages, needing_ocr, ruled)

    # La voie se décide avant de produire quoi que ce soit : calculer un rendu
    # par colonnes pour le jeter ensuite parce que l'arbre l'emporte, c'est du
    # travail payé pour rien.
    by_tree = False
    with _step(timings, "arbre"):
        tree = inspect(payload)

    if tree.usable:
        # Le document porte son plan : il donne les cellules d'un tableau sans
        # qu'on ait à les deviner, et désigne son propre décor comme
        # « Artifact » au lieu de nous le faire compter.
        with _step(timings, "arbre"):
            structured = tagged_reader.read(
                payload, {key: " ".join(parts) for key, parts in marks.items()}
            )
            if any(body.strip() for body in structured):
                pages, by_tree = structured, True

    if not by_tree and not ruled and any(pages_words):
        # Sans arbre et sans filets, c'est la disposition qui fait foi.
        # pdfplumber voit sinon des tableaux là où il n'y a que des barres de
        # fraction : sur les 99 pages du programme national, aucun filet
        # n'atteint la moitié de la largeur de page.
        with _step(timings, "colonnes"):
            try:
                pages = _read_columns(pages_words, widths)
            except Exception:  # noqa: BLE001
                logger.warning("lecture par colonnes impossible, filets conservés")

    described = PdfProfile(
        pages=len(coverages),
        tagged=is_tagged(payload),
        text_coverage=round(median(coverages), 4) if coverages else 0.0,
        pages_needing_ocr=needing_ocr,
    )

    # Rattrapage mot à mot : filet de secours, jamais routine. Il ignore les
    # polices, donc il confond un « ² » légitime avec un symbole Symbol et
    # écrit « ax≤ » à la place de « ax² ». On ne l'appelle que si la
    # correction des polices n'a rien pu faire et que le texte reste mauvais.
    plan = RepairPlan({}, {}, [])
    if repair and not any(font.changed for font in fonts):
        with _step(timings, "rattrapage"):
            if assess(assemble(pages)).score < _RESCUE_FLOOR:
                plan = build_plan(pages_words)
                if not plan.is_empty:
                    pages = [apply_plan(page, plan) for page in pages]
                    logger.info("rattrapage mot à mot", extra={"words": len(plan.words)})

    return PdfReading(pages, plan, fonts, described, ruled, script, timings, tree, by_tree)


def read_pdf_pages(payload: bytes) -> list:
    """Texte de chaque page, dans l'ordre. Une page vide reste une chaîne vide."""

    return read_pdf_document(payload).pages


def assemble(pages: list) -> str:
    """Recomposer un document à partir de ses pages, avec les repères."""

    return normalise(
        "\n\n".join(
            f"## p. {number}\n\n{body.strip()}"
            for number, body in enumerate(pages, start=1)
            if body and body.strip()
        )
    )


def needs_ocr(text: str, page_count: int, minimum_per_page: int) -> bool:
    """Un PDF de scans n'a pas de couche texte, ou presque pas."""

    if page_count <= 0:
        return False
    return len(text.strip()) < minimum_per_page * page_count


def load_pdf(payload: bytes, **_ignored) -> str:
    """Lecture simple d'un PDF, sans réparation.

    La route `/extract` utilise `read_pdf_pages` puis répare les pages
    illisibles. Cette fonction reste pour les usages hors ligne : scripts,
    tests, exploration.
    """

    return clean_ocr_text(assemble(read_pdf_document(payload).pages))


def load_docx(payload: bytes) -> str:
    try:
        import docx
    except ImportError as error:  # pragma: no cover
        raise UnsupportedMediaType("python-docx is not installed") from error

    document = docx.Document(io.BytesIO(payload))
    blocks = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(character for character in style if character.isdigit()) or "1"
            blocks.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return normalise("\n\n".join(blocks))


_LOADERS: Dict[str, Callable[[bytes], str]] = {
    "text/plain": load_plain_text,
    "text/markdown": load_plain_text,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": load_docx,
}


def load(payload: bytes, media_type: str, **pdf_options) -> str:
    if media_type == "application/pdf":
        return load_pdf(payload, **pdf_options)
    loader = _LOADERS.get(media_type)
    if loader is None:
        raise UnsupportedMediaType(f"type non supporté : {media_type}")
    return loader(payload)


# --------------------------------------------------------------------------- sections

_SECTION_SPLIT = re.compile(r"\n\s*\n")
_HEADING_LINE = re.compile(r"^\s{0,3}#{1,6}\s+(?P<title>.*\S)\s*$")


def to_sections(text: str) -> list:
    """Découper le texte extrait en blocs, en gardant le chemin des titres.

    Le repère sert aux citations : un élève doit pouvoir être renvoyé à
    l'endroit exact du document source, pas à un fragment anonyme.
    """

    sections = []
    heading_path: list = []
    position = 0

    for raw in _SECTION_SPLIT.split(text):
        block = raw.strip()
        if not block:
            continue

        body_lines = []
        for line in block.split("\n"):
            heading = _HEADING_LINE.match(line)
            if heading:
                level = len(line) - len(line.lstrip("# \t"))
                level = max(1, line.strip().count("#", 0, 6) or 1)
                del heading_path[level - 1 :]
                heading_path.append(heading.group("title").strip())
                continue
            body_lines.append(line)

        body = "\n".join(body_lines).strip()
        if not body:
            continue

        locator = " > ".join(heading_path) if heading_path else f"§{position + 1}"
        sections.append(
            {
                "position": position,
                "locator": locator,
                "text": body,
                "characters": len(body),
            }
        )
        position += 1

    return sections


# --------------------------------------------------------------------- nettoyage

_PAGE_MARKER = re.compile(r"^## p\. \d+$", re.MULTILINE)
_HYPHENATED = re.compile(r"(\w)-\n(\w)", re.UNICODE)
_LIST_ITEM = re.compile(r"^\s*([-•*—]|\d+[.)])\s+")
_ENDS_SENTENCE = re.compile(r"[.!?…:;»\"]\s*$")


def _repeated_lines(pages: list, threshold: float = 0.04) -> set:
    """Repérer les en-têtes et pieds de page.

    Une ligne qui revient sur beaucoup de pages est un habillage, pas du
    contenu. Le seuil est très bas — 4 % — parce qu'un pied de page change
    avec la partie du document : le programme national en porte onze
    variantes (« Seconde S », « Premières S1 et S3 », « Terminales »,
    « Séries L », plus les versions paire et impaire où le numéro passe de la
    fin au début), chacune sur 8 à 12 pages seulement. À 20 %, aucune
    n'atteignait le seuil et les 76 pieds de page restaient dans le corps.

    Ce seuil bas ne peut pas emporter du contenu, car on ne regarde que les
    deux premières et trois dernières lignes de chaque page : un habillage s'y
    trouve toujours, un paragraphe rarement. Les lignes de tableau sont
    exclues d'office.
    """

    if len(pages) < 3:
        return set()

    counts: Dict[str, int] = {}
    for page in pages:
        lines = [line.strip() for line in page.split("\n") if line.strip()]
        # Un pied de page n'est pas toujours en bas de la feuille — dans le
        # programme national il est à 75 % de la hauteur — mais il est
        # toujours en tête ou en queue du texte de sa page.
        edges = set(lines[:2] + lines[-3:])
        for line in edges:
            # Les lignes de tableau sont du contenu, jamais de l'habillage :
            # les exclure évite de supprimer une ligne de programme répétée.
            if line.startswith("|"):
                continue
            # Sans chiffre, ce n'est pas un titre courant. « INTRODUCTION
            # GENERALE » revient en tête de quatre parties du programme
            # national, « DISCIPLINE : Mathématiques » ouvre chaque bloc de
            # niveau de PHARES : les deux étaient effacés — de la perte de
            # contenu. Un vrai pied de page porte presque toujours un numéro
            # de page ou une année, et c'est lui qui varie.
            if not any(character.isdigit() for character in line):
                continue
            if 3 <= len(line) <= 120:
                counts[_normalise_digits(line)] = counts.get(_normalise_digits(line), 0) + 1

    minimum = max(3, int(len(pages) * threshold))
    return {line for line, count in counts.items() if count >= minimum}


_DIGIT_RUN = re.compile(r"\d+")


def _normalise_digits(line: str) -> str:
    """« page 9 » et « page 12 » sont le même habillage.

    On réduit chaque **suite** de chiffres à un seul marqueur, pas chaque
    chiffre : remplacer caractère par caractère laisse « 2006 1 » et
    « 2006 12 » différents, et le pied de page n'est alors jamais reconnu
    comme répété — chaque page en porte une variante unique.

    Les espaces se normalisent aussi : « Année 200 6 » — l'espace parasite
    vient de l'extraction — doit rejoindre « Année 2006 ». Chiffres d'abord,
    puis fusion des marqueurs voisins, puis espaces.
    """

    reduced = _DIGIT_RUN.sub("#", line)
    reduced = re.sub(r"#(?:\s*#)+", "#", reduced)
    return re.sub(r"\s+", " ", reduced).strip()


def _join_broken_lines(text: str) -> str:
    """Recoller les lignes coupées au milieu d'une phrase.

    Un OCR retourne à la ligne à chaque ligne physique du scan. On ne recolle
    que lorsque la ligne suivante continue visiblement la phrase : titres,
    listes et fins de phrase sont laissés intacts.
    """

    lines = text.split("\n")
    output: List[str] = []

    for line in lines:
        stripped = line.strip()
        if not output:
            output.append(stripped)
            continue

        previous = output[-1]
        continues = (
            previous
            and stripped
            and not _ENDS_SENTENCE.search(previous)
            and not previous.startswith("#")
            and not stripped.startswith("#")
            and not _LIST_ITEM.match(stripped)
            and not _LIST_ITEM.match(previous)
            and stripped[0].islower()
        )
        if continues:
            output[-1] = f"{previous} {stripped}"
        else:
            output.append(stripped)

    return "\n".join(output)


def clean_ocr_text(text: str) -> str:
    """Nettoyage déterministe du texte extrait.

    Corrige ce qui est mécanique : césures, lignes coupées, habillage de page.
    Ne corrige pas les colonnes entrelacées, qui relèvent de l'analyse de mise
    en page — pas de règles.
    """

    if not text.strip():
        return text

    # Césures : "mathéma-\ntiques" -> "mathématiques"
    text = _HYPHENATED.sub(r"\1\2", text)

    # Découpage par page pour repérer l'habillage répété.
    markers = list(_PAGE_MARKER.finditer(text))
    if markers:
        pages = []
        for index, marker in enumerate(markers):
            start = marker.end()
            end = markers[index + 1].start() if index + 1 < len(markers) else len(text)
            pages.append(text[start:end])
        noise = _repeated_lines(pages)
        if noise:
            kept = []
            for line in text.split("\n"):
                if _normalise_digits(line.strip()) in noise:
                    continue
                kept.append(line)
            text = "\n".join(kept)

    text = _join_broken_lines(text)

    # Recoller les mots que l'extraction a coupés — « théor ème »,
    # « déc embre » — sur preuve interne : la forme entière existe ailleurs
    # dans le document. Voir app/core/lexicon.py.
    text = lexicon.mend(text)
    return normalise(text)
