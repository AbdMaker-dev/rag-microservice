from __future__ import annotations

import pytest

from app.core.extraction import UnsupportedMediaType, load, normalise, to_sections


def test_les_lignes_vides_survivent():
    # Elles servent de frontière de bloc en aval.
    assert normalise("a\n\n\n\n\nb") == "a\n\nb"


def test_retours_chariot_et_caracteres_de_controle_retires():
    assert normalise("a\r\nb\x00c") == "a\nbc"


def test_type_non_supporte_refuse():
    with pytest.raises(UnsupportedMediaType):
        load(b"...", "application/zip")


def test_markdown_traite_comme_du_texte():
    assert load("# Titre\n\ncorps".encode("utf-8"), "text/markdown") == "# Titre\n\ncorps"


def test_sections_portent_le_chemin_des_titres():
    sections = to_sections("# Maths\n\n## Algèbre\n\nune équation\n\n# Français\n\nun texte")
    assert [s["locator"] for s in sections] == ["Maths > Algèbre", "Français"]


def test_sections_numerotees_sans_titre():
    sections = to_sections("premier bloc\n\ndeuxième bloc")
    assert [s["position"] for s in sections] == [0, 1]
    assert sections[0]["locator"] == "§1"


def test_texte_vide_ne_produit_aucune_section():
    assert to_sections("   \n\n  ") == []


# ----------------------------------------------------------------- qualité

from app.core.quality import assess, corrupted_pages  # noqa: E402


def test_texte_francais_correct_bien_note():
    q = assess("Les compétences exigibles portent sur la géométrie plane.")
    assert q.score > 0.95


def test_texte_a_encodage_casse_mal_note():
    q = assess("Les compŽtences portent sur la gŽomŽtrie et les probl(cid:143)mes.")
    assert q.score < 0.8
    assert q.cid_markers == 1


def test_formules_mathematiques_ne_penalisent_pas():
    # Les variables d'une lettre et les symboles usuels sont attendus.
    q = assess("Soit f dérivable sur I ; si f'(x) ≥ 0 alors f croît. π ≈ 3,14")
    assert q.score > 0.85


def test_tableau_markdown_bien_note():
    q = assess("| Semaine 2 | Activité Numérique | Les nombres décimaux |")
    assert q.score > 0.9


def test_texte_vide_note_a_zero():
    assert assess("").score == 0.0


def test_pages_corrompues_reperees():
    pages = [
        "Une page parfaitement lisible en français correct.",
        "Une page avec (cid:143) des marqueurs (cid:12) d'échec.",
        "Encore une page saine et bien écrite en français.",
    ]
    assert corrupted_pages(pages) == [2]


def test_page_vide_ignoree():
    assert corrupted_pages(["texte correct en français", "   ", ""]) == []


# ------------------------------------------------------------------ index

from app.db.repository import to_vector_literal  # noqa: E402


def test_vecteur_au_format_pgvector():
    assert to_vector_literal([1.0, -0.5]) == "[1.0,-0.5]"


def test_vecteur_vide_refuse():
    with pytest.raises(ValueError):
        to_vector_literal([])


# ------------------------------------------------- contrat de la réponse

from app.models.schemas import ExtractionQuality, ExtractResponse  # noqa: E402


def test_reponse_extraction_se_construit():
    """Le modèle refuse tout champ inconnu ; un nom mal orthographié dans la
    route ne se voyait qu'au premier appel réel."""

    reponse = ExtractResponse(
        request_id="r1",
        filename="doc.pdf",
        media_type="application/pdf",
        text="Un texte correct en français.",
        characters=29,
        sections=[],
        quality=ExtractionQuality(
            score=0.95, word_plausibility=1.0, cid_markers=0, words_repaired=12
        ),
    )
    assert reponse.quality.words_repaired == 12
    assert reponse.model_dump(by_alias=True)["quality"]["wordsRepaired"] == 12


def test_la_reponse_porte_l_analyse_du_document_sans_casser_le_contrat():
    """`analysis` est additive : un appelant qui l'ignore voit la même réponse."""

    from app.models.schemas import (
        DocumentAnalysis,
        ExtractionQuality,
        ExtractResponse,
        FontDiagnosis,
    )

    reponse = ExtractResponse(
        request_id="r-1",
        filename="programme.pdf",
        media_type="application/pdf",
        text="Compétences exigibles",
        characters=21,
        sections=[],
        quality=ExtractionQuality(
            score=0.99, word_plausibility=0.99, cid_markers=0, words_repaired=0
        ),
        analysis=DocumentAnalysis(
            tagged=False,
            text_coverage=0.149,
            pages_needing_ocr=[3],
            fonts=[FontDiagnosis(font="Times-Roman", table="mac_roman",
                                 confidence=1.0, samples=1436)],
        ),
    )

    rendu = reponse.model_dump(by_alias=True)
    assert rendu["analysis"]["pagesNeedingOcr"] == [3]
    assert rendu["analysis"]["fonts"][0]["table"] == "mac_roman"
    assert rendu["analysis"]["textCoverage"] == 0.149
    # Les clés historiques restent en place et au même endroit.
    assert rendu["contractVersion"] == "1.0"
    assert rendu["quality"]["score"] == 0.99


def test_le_type_reel_du_fichier_prime_sur_le_type_declare():
    """Un navigateur devine l'extension et se trompe régulièrement.

    Suivre la signature ferme deux problèmes : le dépôt légitime mal étiqueté,
    et le fichier arbitraire annoncé comme PDF qui entrerait dans un analyseur
    qui ne l'attend pas.
    """

    from app.core.extraction import sniff

    assert sniff(b"%PDF-1.4\n%...") == "application/pdf"
    assert sniff(b"   %PDF-1.7") == "application/pdf"
    assert sniff(b"PK\x03\x04\x14\x00\x06\x00").endswith("wordprocessingml.document")
    # Pas de signature reconnue : on ne refuse pas pour autant, un fichier
    # texte n'en a pas.
    assert sniff(b"Bonjour, ceci est du texte.") is None
    assert sniff(b"") is None


def test_un_pied_de_page_qui_varie_avec_la_partie_du_document_est_retire():
    """Le programme national porte onze variantes de pied de page.

    « Seconde S », « Premières S1 et S3 », « Terminales », « Séries L », plus
    les versions paire et impaire où le numéro passe de la fin au début —
    chacune sur 8 à 12 pages seulement. À 20 %, aucune n'atteignait le seuil
    et 76 pieds de page restaient dans le corps du texte.
    """

    from app.core.extraction import _repeated_lines

    # Le contenu doit être distinct autrement que par un chiffre : la
    # normalisation réduit toute suite de chiffres à un marqueur, si bien que
    # « paragraphe 3 » et « paragraphe 7 » sont le même texte à ses yeux.
    mots = ["vecteurs", "barycentre", "trigonométrie", "polynômes", "suites",
            "dérivées", "intégrales", "probabilités", "statistiques", "limites",
            "complexes", "arithmétique"]
    pages = []
    for numero in range(1, 25):
        partie = "Seconde S" if numero <= 12 else "Terminales S1 et S3"
        pages.append(
            f"Chapitre sur les {mots[numero % len(mots)]} et leurs propriétés.\n"
            f"Programmes de mathématiques - {partie} - Année 2006 {numero}"
        )

    bruit = _repeated_lines(pages)

    assert any("Seconde S" in ligne for ligne in bruit)
    assert any("Terminales" in ligne for ligne in bruit)
    # Le contenu propre à chaque page ne doit jamais être confondu avec de
    # l'habillage, même avec un seuil aussi bas.
    assert not any("Chapitre sur les" in ligne for ligne in bruit)


def test_un_titre_repete_sans_chiffre_n_est_pas_de_l_habillage():
    """« INTRODUCTION GENERALE » ouvre quatre parties du programme national.

    Le filtre l'effaçait — de la perte de contenu. Un vrai pied de page porte
    presque toujours un numéro de page ou une année ; un titre de partie, non.
    """

    from app.core.extraction import _repeated_lines

    pages = []
    for numero in range(1, 25):
        tete = "INTRODUCTION GENERALE" if numero % 6 == 0 else f"Chapitre {numero}"
        pages.append(f"{tete}\nDu contenu qui change.\nHarmonisation Page {numero}")

    bruit = _repeated_lines(pages)

    assert not any("INTRODUCTION" in ligne for ligne in bruit)
    assert any("Harmonisation" in ligne for ligne in bruit)


def test_un_pied_de_page_aux_espaces_parasites_est_quand_meme_reconnu():
    """« Année 200 6 71 » : l'espace dans « 200 6 » cassait la normalisation."""

    from app.core.extraction import _normalise_digits

    assert _normalise_digits("Année 200 6 71") == _normalise_digits("Année 2006 12")
    assert _normalise_digits("Seconde  S -  Année 2006 5") == _normalise_digits(
        "Seconde S - Année 2006 71"
    )


def test_une_equation_word_ne_se_perd_jamais():
    """« Résoudre ax²+bx+c=0 » ressortait « Résoudre  » — une perte muette.

    Les équations Word vivent dans l'espace de noms OMML, que python-docx
    ignore. Sur une plateforme de mathématiques, c'est le pire mode d'échec :
    le texte a l'air complet.
    """

    import io

    import docx
    import docx.oxml

    from app.core.extraction import load_docx

    document = docx.Document()
    paragraph = document.add_paragraph("Résoudre ")
    paragraph._element.append(docx.oxml.parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        "<m:r><m:t>ax²+bx+c=0</m:t></m:r></m:oMath>"
    ))
    document.add_paragraph("pour tout x réel.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = load_docx(buffer.getvalue())

    assert "Résoudre ax²+bx+c=0" in text


def test_le_word_97_2003_se_reconnait_a_ses_octets():
    from app.core.extraction import is_legacy_doc

    assert is_legacy_doc(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 8)
    assert not is_legacy_doc(b"PK\x03\x04")
    assert not is_legacy_doc(b"%PDF-1.4")


def test_le_contrat_d_indexation_porte_le_cours_et_la_hierarchie_du_pays():
    """Le professeur crée son cours AVANT de déposer des documents.

    Chaque document indexé référence ce cours — c'est ce qui permettra à la
    génération de retrouver les documents d'UN cours, pas d'un périmètre. Et
    le périmètre parle la langue du pays : niveau (primaire, CEM, lycée),
    série quand elle existe, classe.
    """

    from app.models.schemas import IndexRequest, Scope

    requete = IndexRequest(
        request_id="r-1",
        document_id="doc-42",
        course_id="cours-7",
        title="Produit scalaire",
        scope=Scope(
            country="SN", subject="maths", level="lycee", track="S",
            grade="seconde", curriculum_version="2006",
        ),
        text="Le produit scalaire de deux vecteurs.",
    )

    rendu = requete.model_dump(by_alias=True)
    assert rendu["courseId"] == "cours-7"
    assert rendu["scope"]["level"] == "lycee"
    assert rendu["scope"]["track"] == "S"

    # Le primaire et le CEM n'ont pas de série : le champ reste simplement vide.
    primaire = Scope(country="SN", subject="maths", level="primaire",
                     grade="CM2", curriculum_version="2006")
    assert primaire.track == ""


def test_un_support_de_cours_exige_son_cours_et_un_programme_n_en_a_pas():
    """Les deux erreurs inverses sont refusées.

    Un support sans cours serait introuvable à la génération ; un programme
    officiel rattaché à un cours cesserait d'être la référence commune du
    périmètre.
    """

    import base64

    from fastapi.testclient import TestClient

    from app.main import create_app

    client = TestClient(create_app())
    # Le secret vient du conftest : le test ne dépend d'aucun environnement.
    token = {"X-Service-Token": "test-secret-value-of-at-least-32-chars"}
    base = {
        "requestId": "r-1",
        "documentId": "doc-1",
        "title": "Essai",
        "scope": {
            "country": "SN", "subject": "maths", "level": "lycee",
            "track": "S", "grade": "seconde", "curriculumVersion": "2006",
        },
        "text": "Le produit scalaire de deux vecteurs du plan.",
    }

    orphelin = client.post("/index", headers=token, json={**base, "role": "support-cours"})
    assert orphelin.status_code == 422
    assert orphelin.json()["detail"]["code"] == "COURSE_ID_REQUIRED_FOR_COURSE_MATERIAL"

    attache = client.post(
        "/index", headers=token,
        json={**base, "role": "programme-officiel", "courseId": "cours-7"},
    )
    assert attache.status_code == 422
    assert attache.json()["detail"]["code"] == "OFFICIAL_CURRICULUM_HAS_NO_COURSE"


def test_les_embeddings_partent_par_lots():
    """180 chunks en un appel dépassaient le délai CPU : 503 en production.

    Le lot borne chaque appel ; le délai redevient une garantie par lot.
    """

    import asyncio

    from app.config import get_settings
    from app.core.embeddings import OllamaEmbeddingProvider

    calls: list = []

    class _Client:
        async def post(self, url, json=None, timeout=None):
            calls.append(len(json["input"]))

            class _Response:
                def raise_for_status(self):
                    pass

                def json(self):
                    return {"embeddings": [[0.0] * 4] * len(json["input"])}

            return _Response()

    settings = get_settings()
    provider = OllamaEmbeddingProvider(settings, _Client())
    vectors = asyncio.run(provider.embed([f"chunk {i}" for i in range(37)]))

    assert len(vectors) == 37
    assert all(size <= settings.embedding_batch_size for size in calls)
    assert len(calls) >= 3


def test_la_route_search_transmet_vraiment_ses_filtres():
    """Deux correctifs précédents avaient « ajouté » ces filtres sans les
    brancher : le remplacement de texte visait la mauvaise indentation et
    échouait en silence. Le contrat acceptait courseId et role — et les
    ignorait. Ce test lit le code : si un champ du contrat n'atteint pas la
    recherche, il casse.
    """

    import inspect

    from app.api import routes_search
    from app.core.retrieval import Retriever

    source = inspect.getsource(routes_search.search)
    assert "course_id=body.course_id" in source
    assert "role=body.role" in source

    signature = inspect.signature(Retriever.search)
    assert "course_id" in signature.parameters
    assert "role" in signature.parameters


def test_une_annale_est_commune_au_perimetre_comme_un_programme():
    """Un sujet d'examen passé n'appartient à aucun cours : rattaché à un
    cours, il cesserait d'être la référence commune du périmètre."""

    import base64

    from fastapi.testclient import TestClient

    from app.main import create_app

    client = TestClient(create_app())
    token = {"X-Service-Token": "test-secret-value-of-at-least-32-chars"}
    response = client.post("/index", headers=token, json={
        "requestId": "r-a", "documentId": "bac-2024-maths-s2",
        "title": "BAC 2024 — maths S2", "role": "annale", "courseId": "cours-7",
        "scope": {"country": "SN", "subject": "maths", "level": "secondaire",
                  "track": "S2", "grade": "terminale", "curriculumVersion": "2006"},
        "text": "Exercice 1 : similitudes directes.",
    })
    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "OFFICIAL_CURRICULUM_HAS_NO_COURSE"
